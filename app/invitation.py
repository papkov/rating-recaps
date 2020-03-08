from docx import Document
import zipfile
from datetime import date
from shutil import copyfile
import re


def get_invitation(team_id, invitation_id, position, university, first_name, second_name, surname, team_name, recaps):
    substitution = {
        'position': position,
        'university': university,
        'surname': surname,
        'team_name': f'"{team_name}"',
        'fnamef': second_name[0],
        'namef': first_name[0],
        'fname': second_name,
        'name': first_name,
        'recaps': '\n'.join(recaps),
        'ID': str(invitation_id),
        'date': str(date.today().strftime("%d"))
    }

    template = './data/invitation_template.docx'
    invitation = f'./invitations/{team_id}_Приглашение_ОЧВР_{team_name}.docx'
    # copyfile(template, invitation)
    # return docx_replace(template, invitation, substitution)
    doc = Document(template)
    r = docx_replace_2(doc, substitution)
    doc.save(invitation)
    return r


def docx_replace(old_file, new_file, substitution):
    """
    Replace by dictionary
    https://stackoverflow.com/questions/17850227/text-replace-in-docx-and-save-the-changed-file-with-python-docx
    :param old_file:
    :param new_file:
    :param substitution:
    :return:
    """
    with zipfile.ZipFile(old_file, 'r') as zin:
        with zipfile.ZipFile(new_file, 'w') as zout:
            for item in zin.infolist():
                buffer = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    res = buffer.decode("utf-8")
                    print(res)
                    for r in substitution:
                        print(r, substitution[r])
                        res = res.replace(r, substitution[r])
                    buffer = res.encode("utf-8")
                zout.writestr(item, buffer)

    return True


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def docx_replace_2(doc_obj, substitution):
    for r in substitution:
        regex = re.compile(r"{}".format(r))
        replace = r"{}".format(substitution[r])
        docx_replace_regex(doc_obj, regex, replace)
    return True
